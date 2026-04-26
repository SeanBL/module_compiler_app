from __future__ import annotations

from pathlib import Path
from typing import List, Optional

from docx import Document
from docx.oxml.ns import qn
from docx.document import Document as _Document
from docx.table import Table
from docx.text.paragraph import Paragraph

from ..models.blocks import ParagraphBlock, BulletsBlock, Block
from ..models.raw_models import (
    RawSlide,
    RawEngage1Item,
    RawQuizQuestion,
    RawQuizOption,
)

# --------------------------------------------------
# Utilities
# --------------------------------------------------
def iter_block_items(parent):
    """
    Yield paragraphs and tables in document order.
    """
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    else:
        parent_elm = parent._tc

    for child in parent_elm.iterchildren():
        if child.tag.endswith('}p'):
            yield Paragraph(child, parent)
        elif child.tag.endswith('}tbl'):
            yield Table(child, parent)


def normalize(text: str) -> str:
    if not text:
        return ""
    return " ".join(text.replace("\u00A0", " ").split()).strip()


def is_list_paragraph(p) -> bool:
    try:
        pPr = p._p.pPr
        if pPr is None:
            return False
        return pPr.find(qn("w:numPr")) is not None
    except Exception:
        return False


def canonical_col_label(raw: str) -> str:
    t = normalize(raw).lower().strip()

    if "/" in t:
        t = t.split("/", 1)[0].strip()

    if "notes" in t:
        return "notes"
    if "english" in t:
        return "english"
    if "image" in t:
        return "image"
    if "button" in t:
        return "button_labels"

    return t


def is_slide_header(text: str) -> bool:
    t = text.lower()
    return (
        t.startswith("header:")
        or t.startswith("slide header")
        or t.startswith("slide (header)")
    )

def is_quiz_marker(text: str) -> bool:
    t = normalize(text).lower()
    return (
        t.startswith("quiz_")
        and (
            t.endswith("_inline")
            or t.endswith("_application")
            or t.endswith("_final")
        )
    )

def extract_quiz_scope(marker: str) -> str:
    t = normalize(marker).lower()

    if t.endswith("_final"):
        return "final"

    if t.endswith("_application"):
        return "application"

    return "inline"

def starts_numbered_engage_section(blocks: List[Block]) -> bool:
    if not blocks:
        return False

    first = blocks[0]

    if not isinstance(first, ParagraphBlock):
        return False

    text = first.text.strip()

    # Handles: 1. Text, 2. Text, etc.
    return (
        len(text) >= 3
        and text[0].isdigit()
        and text[1] == "."
        and text[2].isspace()
    )

def extract_module_metadata(doc) -> dict:
    module_title = None
    module_id = None

    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            text = normalize(block.text)

            if text.lower().startswith("module title"):
                parts = text.split(":", 1)
                if len(parts) == 2:
                    module_title = parts[1].strip()

            if text.lower().startswith("module id"):
                parts = text.split(":", 1)
                if len(parts) == 2:
                    module_id = parts[1].strip()

        # Stop early once both found (SAFE optimization)
        if module_title and module_id:
            break

    return {
        "module_title": module_title,
        "module_id": module_id
    }

def is_real_bullet(p):
    try:
        # Case 1: standard numbering (works for some bullets)
        pPr = p._element.pPr
        if pPr is not None and pPr.numPr is not None:
            return True

        # Case 2: style-based bullets (VERY common in tables)
        style_name = p.style.name.lower()
        if "list" in style_name:
            return True

    except Exception:
        pass

    return False

def extract_cell_blocks(cell) -> List[Block]:
    blocks: List[Block] = []
    bullet_items: List[str] = []

    def flush_bullets():
        nonlocal bullet_items
        if bullet_items:
            blocks.append(
                BulletsBlock(
                    type="bullets",
                    items=bullet_items.copy()
                )
            )
            bullet_items = []

    for p in cell.paragraphs:
        text = normalize(p.text)
        if not text:
            flush_bullets()
            continue

        # Detect bullet manually OR fallback pattern
        is_bullet = False

        # Case 1: real bullet character
        if text.startswith("•"):
            is_bullet = True
            text = text.replace("•", "").strip()

        # Case 2: dash bullet (YOUR CURRENT CASE)
        elif text.startswith("- "):
            is_bullet = True
            text = text[2:].strip()   # removes "- "

        # Case 3: Word list formatting
        elif is_real_bullet(p):
            is_bullet = True

        # # Case 4: semantic list (colon trigger)
        # elif (
        #     len(blocks) > 0
        #     and isinstance(blocks[-1], ParagraphBlock)
        #     and blocks[-1].text.strip().endswith(":")
        #     and text[0].isupper()  # basic guard
        #     and not text.endswith(".")  # prevents full sentences from joining bullets
        # ):
        #     is_bullet = True

        # Final decision
        if is_bullet:
            # FIX: prevent intro line from being treated as bullet
            if not bullet_items and text.endswith(":"):
                flush_bullets()
                blocks.append(
                    ParagraphBlock(
                        type="paragraph",
                        text=text
                    )
                )
            else:
                bullet_items.append(text)
        else:
            flush_bullets()
            blocks.append(
                ParagraphBlock(
                    type="paragraph",
                    text=text
                )
            )

    flush_bullets()
    if any("One person with a fever" in (getattr(b, "text", "") or "") for b in blocks):
        print("🔍 DEBUG TARGET BLOCKS:", blocks)
    return blocks

def finalize_engage1_slide(slide: RawSlide, pending_button_labels: List[str]) -> None:
    if not pending_button_labels:
        raise ValueError(
            f"Engage1 slide {slide.slide_id} has no buttons"
        )

    if not slide.body or not any(
        isinstance(block, (ParagraphBlock, BulletsBlock))
        for section in slide.body
        for block in section
    ):
        raise ValueError(
            f"Engage1 slide {slide.slide_id} has no content blocks"
        )

    sections = slide.body
    intro = sections[0]
    groups = sections[1:]

    if len(groups) != len(pending_button_labels):
        raise ValueError(
            f"Engage1 slide {slide.slide_id} has {len(groups)} sections but {len(pending_button_labels)} buttons."
        )

    slide.engage1_items = [
        RawEngage1Item(
            label=label,
            body=group,
            image=None
        )
        for label, group in zip(pending_button_labels, groups)
    ]

    slide.engage1_intro = intro
    slide.engage1_intro_image = getattr(intro[0], "image", None)
    slide.body = None

# --------------------------------------------------
# Structural Extraction
# --------------------------------------------------

def extract_raw_slides(docx_path: Path) -> dict:
    doc = Document(str(docx_path))
    metadata = extract_module_metadata(doc)

    if not metadata["module_title"]:
        raise ValueError("Missing 'Module title:' in Word document")

    if not metadata["module_id"]:
        raise ValueError("Missing 'Module id:' in Word document")

    slides: List[RawSlide] = []
    slide_index = 0

    current_slide: Optional[RawSlide] = None
    col_labels: List[str] = []

    engage1_items: List[RawEngage1Item] = []
    pending_button_labels: List[str] = []

    quiz_questions: List[RawQuizQuestion] = []
    quiz_inline_mode = False
    quiz_inline_questions: List[RawQuizQuestion] = []

    for block in iter_block_items(doc):

        # -----------------------------
        # Paragraph blocks (markers)
        # -----------------------------
        if isinstance(block, Paragraph):
            text = normalize(block.text)

            # Skip metadata rows
            if text.lower().startswith("module title") or text.lower().startswith("module id"):
                continue

            print("PARAGRAPH:", repr(text))

            if is_quiz_marker(text):
                print("QUIZ MARKER DETECTED:", text)

                # Properly finalize previous slide first
                if current_slide is not None:

                    if current_slide.slide_type == "engage1":
                        finalize_engage1_slide(current_slide, pending_button_labels)

                        # 🔥 CRITICAL RESET
                        pending_button_labels = []
                        current_slide.body = None

                    if current_slide.slide_type == "quiz":
                        current_slide.quiz_questions = quiz_questions
                        current_slide.quiz_type = "mcq"
                        quiz_questions = []

                    slides.append(current_slide)

                slide_index += 1

                scope = extract_quiz_scope(text)

                current_slide = RawSlide(
                    slide_id=f"slide_{slide_index:03d}",
                    header=text,
                    slide_type="quiz",
                )

                current_slide.quiz_scope = scope

                quiz_questions = []
                quiz_inline_mode = True
                continue
            continue

        # -----------------------------
        # Table blocks (real content)
        # -----------------------------
        if not isinstance(block, Table):
            continue

        tbl = block

        if len(tbl.rows) < 2:
            continue

        row_idx = 0

        while row_idx < len(tbl.rows):
            row = tbl.rows[row_idx]
            first_cell = normalize(row.cells[0].text)

            # -----------------------------
            # Start new slide
            # -----------------------------
            if is_slide_header(first_cell):

                # finalize previous slide
                if current_slide is not None:

                    if current_slide.slide_type == "engage1":

                        if not pending_button_labels:
                            raise ValueError(
                                f"Engage1 slide {current_slide.slide_id} has no buttons"
                            )

                        if not current_slide.body or not any(
                            isinstance(block, (ParagraphBlock, BulletsBlock))
                            for section in current_slide.body
                            for block in section
                        ):
                            raise ValueError(
                                f"Engage1 slide {current_slide.slide_id} has no content blocks"
                            )

                        sections = current_slide.body

                        intro = sections[0]
                        groups = sections[1:]

                        if len(groups) != len(pending_button_labels):
                            raise ValueError(
                                f"Engage1 slide {current_slide.slide_id} has {len(groups)} sections but {len(pending_button_labels)} buttons."
                            )

                        engage1_items = []

                        for label, group in zip(pending_button_labels, groups):
                            engage1_items.append(
                                RawEngage1Item(
                                    label=label,
                                    body=group,
                                    image=None
                                )
                            )

                        current_slide.engage1_items = engage1_items
                        current_slide.engage1_intro = intro
                        current_slide.engage1_intro_image = getattr(intro[0], "image", None)
                        current_slide.body = None

                    if current_slide.slide_type == "engage2":

                        if not current_slide.body or not any(
                            isinstance(b, (ParagraphBlock, BulletsBlock))
                            for b in current_slide.body
                        ):
                            raise ValueError(
                                f"Engage2 slide {current_slide.slide_id} has no content blocks"
                            )

                        paragraphs = current_slide.body
                        intro_blocks = [paragraphs[0]]
                        layer_blocks = paragraphs[1:]

                        if not layer_blocks:
                            raise ValueError(
                                f"Engage2 slide {current_slide.slide_id} must have at least one reveal layer"
                            )

                        current_slide.engage2_intro = intro_blocks
                        current_slide.engage2_intro_image = getattr(intro_blocks[0], "image", None)
                        current_slide.engage2_layers = layer_blocks
                        current_slide.body = None

                    if current_slide.slide_type == "quiz":
                        current_slide.quiz_questions = quiz_questions
                        current_slide.quiz_type = "mcq"
                        quiz_questions = []

                    slides.append(current_slide)

                    # attach inline quiz slide
                    if quiz_inline_questions:
                        slide_index += 1
                        quiz_slide = RawSlide(
                            slide_id=f"slide_{slide_index:03d}",
                            header=f"{current_slide.header} (Quiz)",
                            slide_type="quiz",
                        )
                        quiz_slide.quiz_questions = quiz_inline_questions
                        quiz_slide.quiz_type = "mcq"
                        slides.append(quiz_slide)

                        quiz_inline_questions = []
                        quiz_inline_mode = False

                slide_index += 1
                engage1_items = []
                pending_button_labels = []

                current_slide = RawSlide(
                    slide_id=f"slide_{slide_index:03d}",
                    header=first_cell,
                    slide_type="panel",
                )

                label_row = tbl.rows[row_idx + 1]
                col_labels = [
                    canonical_col_label(c.text) for c in label_row.cells
                ]

                print("DEBUG COL LABELS:", col_labels)

                row_idx += 2
                continue

            if current_slide is None:
                row_idx += 1
                continue

            # -----------------------------
            # INLINE QUIZ PARSING
            # -----------------------------
            if quiz_inline_mode:

                if not first_cell:
                    row_idx += 1
                    continue

                if len(row.cells) < 2:
                    row_idx += 1
                    continue

                q_lines = [
                    l.strip()
                    for l in row.cells[0].text.splitlines()
                    if l.strip()
                ]

                prompt_lines = []
                options: List[RawQuizOption] = []

                for line in q_lines:
                    if len(line) > 2 and line[1] == "." and line[0].isalpha():
                        options.append(
                            RawQuizOption(
                                id=line[0].upper(),
                                text=line[2:].strip(),
                            )
                        )
                    else:
                        prompt_lines.append(line)

                prompt = " ".join(prompt_lines).strip()

                a_lines = [
                    l.strip()
                    for l in row.cells[1].text.splitlines()
                    if l.strip()
                ]

                correct_option_id = None
                explanation_lines = []

                for line in a_lines:
                    if line.lower().startswith("answer"):
                        parts = line.split(":")
                        if len(parts) == 2:
                            correct_option_id = parts[1].strip()
                    else:
                        explanation_lines.append(line)

                if correct_option_id:

                    # ----------------------------------------
                    # TRUE / FALSE AUTO-CONVERSION PATCH
                    # ----------------------------------------
                    if not options and correct_option_id.lower() in ("true", "false"):

                        options = [
                            RawQuizOption(id="A", text="True"),
                            RawQuizOption(id="B", text="False"),
                        ]

                        correct_option_id = (
                            "A" if correct_option_id.lower() == "true" else "B"
                        )

                    quiz_questions.append(
                        RawQuizQuestion(
                            prompt=prompt,
                            options=options,
                            correct_option_id=correct_option_id,
                            explanation=" ".join(explanation_lines).strip()
                            if explanation_lines
                            else None,
                        )
                    )

                row_idx += 1
                continue

            # -----------------------------
            # Normal row parsing
            # -----------------------------
            row_data = {}

            for idx, cell in enumerate(row.cells):
                if idx >= len(col_labels):
                    continue
                label = col_labels[idx]
                if not label:
                    continue

                # ✅ ONLY convert english to structured blocks
                if label == "english":
                    cell_blocks = extract_cell_blocks(cell)
                    if cell_blocks:
                        existing = row_data.get(label, [])
                        existing.extend(cell_blocks)
                        row_data[label] = existing
                else:
                    # ✅ Keep everything else as plain text (UNCHANGED)
                    texts = [
                        normalize(p.text)
                        for p in cell.paragraphs
                        if normalize(p.text)
                    ]
                    if texts:
                        row_data[label] = texts

            notes = row_data.get("notes", [])
            english = row_data.get("english", [])
            image = row_data.get("image", [])

            if notes:
                blob = " ".join(n.lower() for n in notes)

                if "slide type = quiz" in blob:
                    current_slide.slide_type = "quiz"
                elif "slide type = engage 1" in blob:
                    current_slide.slide_type = "engage1"
                elif "slide type = engage 2" in blob:
                    current_slide.slide_type = "engage2"

            # -----------------------------
            # ENGAGE1 PARSING (Correct Layout Handling)
            # -----------------------------
            if current_slide.slide_type == "engage1" and english:

                row_has_button = any(
                    isinstance(block, ParagraphBlock)
                    and block.text.strip().lower().startswith("[button]")
                    for block in english
                )

                # Button-definition row
                if row_has_button:
                    for block in english:
                        if isinstance(block, ParagraphBlock):
                            text = block.text.strip()
                            if text.strip().lower().startswith("[button]"):
                                label = text[8:].strip()
                                pending_button_labels.append(label)

                # Content row
                else:
                    sections = current_slide.body or []

                    section_blocks = []

                    for block in english:
                        if isinstance(block, ParagraphBlock):
                            if image:
                                block.image = image[0]
                            section_blocks.append(block)
                        else:
                            section_blocks.append(block)

                    has_numbered_sections = any(
                        starts_numbered_engage_section(section)
                        for section in sections[1:]  # skip intro
                    )

                    current_row_starts_numbered_section = starts_numbered_engage_section(section_blocks)

                    if current_row_starts_numbered_section:
                        sections.append(section_blocks)
                    elif has_numbered_sections and len(sections) > 1:
                        sections[-1].extend(section_blocks)
                    else:
                        sections.append(section_blocks)

                    current_slide.body = sections

            # -----------------------------
            # ENGAGE2 PARSING
            # -----------------------------
            if current_slide.slide_type == "engage2" and english:

                blocks = current_slide.body or []

                for block in english:

                    # Handle ParagraphBlock
                    if isinstance(block, ParagraphBlock):
                        text = block.text.strip()

                        if text.lower().startswith("[button]"):
                            button_label = text[8:].strip()
                            current_slide.engage2_button_label = button_label
                            continue

                        if image:
                            block.image = image[0]

                        blocks.append(block)

                    # Handle BulletsBlock
                    else:
                        blocks.append(block)

                current_slide.body = blocks


            if current_slide.slide_type == "panel" and english:
                blocks = current_slide.body or []

                for block in english:
                    if isinstance(block, ParagraphBlock):
                        if image:
                            block.image = image[0]
                        blocks.append(block)
                    else:
                        blocks.append(block)

                current_slide.body = blocks

            if image:
                image_filename = image[0]

                if current_slide.slide_type == "panel":
                    current_slide.image = image_filename

            row_idx += 1

        # ✅ after finishing this table
        if quiz_inline_mode:
            quiz_inline_mode = False

    # finalize last slide
    if current_slide is not None:

        if current_slide.slide_type == "engage1":

            if not pending_button_labels:
                raise ValueError(
                    f"Engage1 slide {current_slide.slide_id} has no buttons"
                )

            if not current_slide.body or not any(
                isinstance(block, (ParagraphBlock, BulletsBlock))
                for section in current_slide.body
                for block in section
            ):
                raise ValueError(
                    f"Engage1 slide {current_slide.slide_id} has no content blocks"
                )

            sections = current_slide.body

            intro = sections[0]
            groups = sections[1:]

            if len(groups) != len(pending_button_labels):
                raise ValueError(
                    f"Engage1 slide {current_slide.slide_id} has {len(groups)} sections but {len(pending_button_labels)} buttons."
                )

            engage1_items = []

            for label, group in zip(pending_button_labels, groups):
                engage1_items.append(
                    RawEngage1Item(
                        label=label,
                        body=group,
                        image=None
                    )
                )

            current_slide.engage1_items = engage1_items
            current_slide.engage1_intro = intro
            current_slide.engage1_intro_image = getattr(intro[0], "image", None)
            current_slide.body = None

        if current_slide.slide_type == "engage2":

            if not current_slide.body or not any(
                isinstance(b, (ParagraphBlock, BulletsBlock))
                for b in current_slide.body
            ):
                raise ValueError(
                    f"Engage2 slide {current_slide.slide_id} has no content blocks"
                )

            paragraphs = current_slide.body
            intro_blocks = [paragraphs[0]]
            layer_blocks = paragraphs[1:]

            if not layer_blocks:
                raise ValueError(
                    f"Engage2 slide {current_slide.slide_id} must have at least one reveal layer"
                )

            current_slide.engage2_intro = intro_blocks
            current_slide.engage2_intro_image = getattr(intro_blocks[0], "image", None)
            current_slide.engage2_layers = layer_blocks
            current_slide.body = None

        if current_slide.slide_type == "quiz":
            current_slide.quiz_questions = quiz_questions
            current_slide.quiz_type = "mcq"

        slides.append(current_slide)

    return {
        "module_title": metadata["module_title"],
        "module_id": metadata["module_id"],
        "slides": slides
    }

    
